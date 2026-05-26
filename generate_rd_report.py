import sqlite3
import os
from datetime import date, datetime, timedelta
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────
# CONFIGURATION
# ─────────────────────────────────────────────

DB_PATH = "data/tmdb_20260512.db"
OUTPUT_PATH = "output/RD_Labour_Hours_Report_FY2526.docx"

PERIOD_START = date(2025, 5, 1)
PERIOD_END   = date(2026, 4, 30)

RD_PROJECTS = [
    'T003-UI-FlowManagement',
    '0202 Luminwave Testing',
    '0141 CallBox Manager G2',
    '8001 General UI',
    'T001-UI-Development',
    '0150-UI-Backend-G2',
    'T002-UI-GTP',
]

# Staff info: join_date, leave_days within May 2025 - Apr 2026
STAFF_INFO = {
    'jason.eng':  {'name': 'Jason Eng',       'join_date': date(2024, 10, 1), 'leave_days': 6.5},
    'jinghao':    {'name': 'Jing Hao Eng',    'join_date': date(2024, 6, 10), 'leave_days': 9.5},
    'kxlim':      {'name': 'Kai Xuan Lim',    'join_date': date(2025, 1, 1),  'leave_days': 9.0},
    'aina':       {'name': 'Aina Maisarah',   'join_date': date(2025, 10, 6), 'leave_days': 10.0},
    'xavier.ong': {'name': 'Xavier Ong',      'join_date': date(2026, 1, 5),  'leave_days': 3.0},
    'pongwx':     {'name': 'Wei Xiang Pong',  'join_date': date(2025, 1, 1),  'leave_days': 10.0},
}

PUBLIC_HOLIDAYS_PER_YEAR = 11
HOURS_PER_DAY = 8
HEADER_COLOR = RGBColor(189, 215, 238)  # Light blue

# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────

def count_weekdays(start, end):
    """Count weekdays (Mon-Fri) between start and end inclusive."""
    count = 0
    current = start
    while current <= end:
        if current.weekday() < 5:
            count += 1
        current += timedelta(days=1)
    return count

def calc_available_hours(user_id):
    info = STAFF_INFO.get(user_id)
    if not info:
        return 0, 0

    effective_start = max(info['join_date'], PERIOD_START)
    total_weekdays_period = count_weekdays(PERIOD_START, PERIOD_END)
    staff_weekdays = count_weekdays(effective_start, PERIOD_END)

    # Prorate public holidays
    prorated_holidays = PUBLIC_HOLIDAYS_PER_YEAR * (staff_weekdays / total_weekdays_period)
    available_days = staff_weekdays - prorated_holidays - info['leave_days']
    available_hours = available_days * HOURS_PER_DAY
    return round(available_days, 2), round(available_hours, 2)

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = '{:02X}{:02X}{:02X}'.format(rgb.red, rgb.green, rgb.blue)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def style_header_row(row):
    for cell in row.cells:
        set_cell_bg(cell, HEADER_COLOR)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
            if not para.runs:
                run = para.add_run(para.text)
                run.bold = True

def bold_row(row):
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(
            "Dallas Roboter Sdn Bhd  |  R&D Labour Hours Capitalisation Report  |  FY May 2025 – Apr 2026  |  Confidential"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

def add_page_number(doc):
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        # Add page number field on right
        para2 = footer.add_paragraph()
        para2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para2.add_run("Page ")
        run.font.size = Pt(8)
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

# ─────────────────────────────────────────────
# DATABASE QUERY
# ─────────────────────────────────────────────

def fetch_data():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    placeholders = ",".join("?" for _ in RD_PROJECTS)
    query = f"""
        SELECT
            c.ContributeInfoId,
            c.StartDateTime,
            c.ClockInHours,
            c.ProjectTitle,
            c.TopicTitle,
            c.TaskTitle,
            c.SubtaskTitle,
            c.Remarks,
            c.UserId,
            u.FirstName,
            u.LastName
        FROM Contributions c
        JOIN Users u ON c.UserId = u.UserId
        WHERE c.StartDateTime BETWEEN '2025-05-01' AND '2026-04-30'
          AND c.IsValid = 1
          AND c.ProjectTitle IN ({placeholders})
        ORDER BY c.StartDateTime ASC
    """
    rows = conn.execute(query, tuple(RD_PROJECTS)).fetchall()
    conn.close()
    return rows

# ─────────────────────────────────────────────
# BUILD SUMMARIES
# ─────────────────────────────────────────────

def build_summaries(rows):
    by_project = {}
    by_person   = {}
    by_month    = {}
    detail      = []

    for r in rows:
        dt = datetime.strptime(r["StartDateTime"][:10], "%Y-%m-%d")
        staff_name = f"{r['FirstName']} {r['LastName']}"
        hours = r["ClockInHours"] or 0
        month_key = dt.strftime("%b %Y")
        project = r["ProjectTitle"] or ""
        user_id = r["UserId"]

        # By project
        if project not in by_project:
            by_project[project] = {"hours": 0, "staff": set()}
        by_project[project]["hours"] += hours
        by_project[project]["staff"].add(user_id)

        # By person
        if user_id not in by_person:
            by_person[user_id] = {"name": staff_name, "hours": 0, "projects": set()}
        by_person[user_id]["hours"] += hours
        by_person[user_id]["projects"].add(project)

        # By month
        if month_key not in by_month:
            by_month[month_key] = 0
        by_month[month_key] += hours

        # Detail
        detail.append({
            "date": dt.strftime("%d/%m/%Y"),
            "staff": staff_name,
            "project": project,
            "topic": r["TopicTitle"] or "",
            "task": r["TaskTitle"] or "",
            "subtask": r["SubtaskTitle"] or "",
            "remarks": r["Remarks"] or "",
            "hours": hours,
            "user_id": user_id,
        })

    # Build full month list May 2025 - Apr 2026
    months_ordered = []
    current = date(2025, 5, 1)
    while current <= date(2026, 4, 30):
        months_ordered.append(current.strftime("%b %Y"))
        if current.month == 12:
            current = current.replace(year=current.year + 1, month=1)
        else:
            current = current.replace(month=current.month + 1)

    return by_project, by_person, by_month, months_ordered, detail

# ─────────────────────────────────────────────
# BUILD WORD DOCUMENT
# ─────────────────────────────────────────────

def build_report(by_project, by_person, by_month, months_ordered, detail):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── COVER PAGE ──────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DALLAS ROBOTER SDN BHD")
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("R&D Labour Hours Capitalisation Report")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 73, 125)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Financial Year: May 2025 – April 2026")
    run.font.size = Pt(14)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Prepared pursuant to:")
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Section 34(7) Income Tax Act 1967 (Malaysia)")
    run.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MFRS 138 Intangible Assets")
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Report Generated: {date.today().strftime('%d %B %Y')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    # ── SECTION 1: BACKGROUND ────────────────────
    doc.add_heading("1. Background & Basis of Preparation", level=1)
    doc.add_paragraph(
        "This report documents the research and development (R&D) labour hours incurred by "
        "Dallas Roboter Sdn Bhd for the financial year covering May 2025 to April 2026. "
        "The hours recorded herein represent qualifying R&D activities undertaken by the Company's "
        "engineering and application development teams in the course of developing proprietary "
        "software systems, automation solutions, and technology platforms."
    )
    doc.add_paragraph(
        "This report has been prepared in accordance with Malaysian Financial Reporting Standard "
        "MFRS 138 (Intangible Assets), which governs the recognition and measurement of internally "
        "generated intangible assets, and Section 34(7) of the Income Tax Act 1967 (Malaysia), "
        "which provides for double deduction of qualifying R&D expenditure. Labour hours have been "
        "extracted directly from the Company's Time Management System (TMS) and represent actual "
        "time logged by staff against eligible R&D projects."
    )

    # ── SECTION 2: BY PROJECT ────────────────────
    doc.add_heading("2. Summary by Project", level=1)
    doc.add_paragraph(
        "The following table summarises total R&D labour hours logged per project during the reporting period."
    )

    headers = ["No.", "Project Title", "No. of Staff", "Total Hours"]
    col_widths = [Inches(0.4), Inches(3.5), Inches(1.2), Inches(1.2)]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        hdr.cells[i].width = w
        hdr.cells[i].text = h
    style_header_row(hdr)

    total_hours = 0
    for idx, (proj, data) in enumerate(sorted(by_project.items()), 1):
        row = table.add_row()
        row.cells[0].text = str(idx)
        row.cells[1].text = proj
        row.cells[2].text = str(len(data["staff"]))
        row.cells[3].text = f"{data['hours']:.2f}"
        total_hours += data["hours"]

    total_row = table.add_row()
    total_row.cells[0].text = ""
    total_row.cells[1].text = "TOTAL"
    total_row.cells[2].text = ""
    total_row.cells[3].text = f"{total_hours:.2f}"
    bold_row(total_row)

    # ── SECTION 3: BY STAFF ──────────────────────
    doc.add_heading("3. Summary by Staff", level=1)
    doc.add_paragraph(
        "The following table summarises total R&D labour hours contributed by each staff member."
    )

    headers = ["No.", "Staff Name", "No. of Projects", "Total Hours"]
    col_widths = [Inches(0.4), Inches(2.8), Inches(1.5), Inches(1.2)]
    table2 = doc.add_table(rows=1, cols=len(headers))
    table2.style = "Table Grid"
    hdr2 = table2.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        hdr2.cells[i].width = w
        hdr2.cells[i].text = h
    style_header_row(hdr2)

    total_staff_hours = 0
    for idx, (uid, data) in enumerate(sorted(by_person.items(), key=lambda x: -x[1]["hours"]), 1):
        row = table2.add_row()
        row.cells[0].text = str(idx)
        row.cells[1].text = data["name"]
        row.cells[2].text = str(len(data["projects"]))
        row.cells[3].text = f"{data['hours']:.2f}"
        total_staff_hours += data["hours"]

    total_row2 = table2.add_row()
    total_row2.cells[0].text = ""
    total_row2.cells[1].text = "TOTAL"
    total_row2.cells[2].text = ""
    total_row2.cells[3].text = f"{total_staff_hours:.2f}"
    bold_row(total_row2)

    # ── SECTION 3.5: UTILIZATION ─────────────────
    doc.add_heading("3.5  Staff R&D Utilization", level=1)
    doc.add_paragraph(
        "The following table presents each staff member's R&D utilization rate for the reporting period. "
        "Available hours are calculated based on the staff member's join date, a 5-day work week at "
        f"{HOURS_PER_DAY} hours per day, less {PUBLIC_HOLIDAYS_PER_YEAR} public holidays (prorated) "
        "and approved leave days."
    )

    headers = ["No.", "Staff Name", "Join Date", "Avail. Days", "Avail. Hours", "R&D Hours", "Utilization (%)"]
    col_widths = [Inches(0.4), Inches(2.0), Inches(0.9), Inches(0.85), Inches(0.9), Inches(0.85), Inches(1.0)]
    table3 = doc.add_table(rows=1, cols=len(headers))
    table3.style = "Table Grid"
    hdr3 = table3.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        hdr3.cells[i].width = w
        hdr3.cells[i].text = h
    style_header_row(hdr3)

    total_avail_hours = 0
    total_rd_hours_util = 0
    util_list = []

    for idx, (uid, data) in enumerate(sorted(by_person.items(), key=lambda x: -x[1]["hours"]), 1):
        avail_days, avail_hours = calc_available_hours(uid)
        rd_hours = data["hours"]
        util_pct = (rd_hours / avail_hours * 100) if avail_hours > 0 else 0
        join_date = STAFF_INFO.get(uid, {}).get("join_date", "")
        join_str = join_date.strftime("%d %b %Y") if isinstance(join_date, date) else "-"

        row = table3.add_row()
        row.cells[0].text = str(idx)
        row.cells[1].text = data["name"]
        row.cells[2].text = join_str
        row.cells[3].text = f"{avail_days:.1f}"
        row.cells[4].text = f"{avail_hours:.2f}"
        row.cells[5].text = f"{rd_hours:.2f}"
        row.cells[6].text = f"{util_pct:.2f}%"

        total_avail_hours += avail_hours
        total_rd_hours_util += rd_hours
        util_list.append(util_pct)

    avg_util = sum(util_list) / len(util_list) if util_list else 0
    total_row3 = table3.add_row()
    total_row3.cells[0].text = ""
    total_row3.cells[1].text = "TOTAL / AVERAGE"
    total_row3.cells[2].text = ""
    total_row3.cells[3].text = ""
    total_row3.cells[4].text = f"{total_avail_hours:.2f}"
    total_row3.cells[5].text = f"{total_rd_hours_util:.2f}"
    total_row3.cells[6].text = f"{avg_util:.2f}%"
    bold_row(total_row3)

    # ── SECTION 4: MONTHLY TREND ─────────────────
    doc.add_heading("4. Monthly Trend", level=1)
    doc.add_paragraph(
        "The following table shows total R&D labour hours logged per month across the reporting period."
    )

    headers = ["Month", "Total Hours"]
    col_widths = [Inches(2.5), Inches(2.0)]
    table4 = doc.add_table(rows=1, cols=len(headers))
    table4.style = "Table Grid"
    hdr4 = table4.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        hdr4.cells[i].width = w
        hdr4.cells[i].text = h
    style_header_row(hdr4)

    monthly_total = 0
    for m in months_ordered:
        hrs = by_month.get(m, 0)
        row = table4.add_row()
        row.cells[0].text = m
        row.cells[1].text = f"{hrs:.2f}"
        monthly_total += hrs

    total_row4 = table4.add_row()
    total_row4.cells[0].text = "TOTAL"
    total_row4.cells[1].text = f"{monthly_total:.2f}"
    bold_row(total_row4)

    # ── SECTION 5: DETAIL LOG ────────────────────
    doc.add_page_break()
    doc.add_heading("5. Detailed Contribution Log", level=1)
    doc.add_paragraph(
        "The following table presents all individual R&D time log entries for the reporting period, "
        "sorted chronologically."
    )

    headers = ["No.", "Date", "Staff", "Project", "Topic", "Task", "Subtask", "Remarks", "Hours"]
    col_widths = [
        Inches(0.3), Inches(0.8), Inches(0.9), Inches(1.2),
        Inches(0.8), Inches(0.8), Inches(0.7), Inches(1.3), Inches(0.5)
    ]
    table5 = doc.add_table(rows=1, cols=len(headers))
    table5.style = "Table Grid"
    hdr5 = table5.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        hdr5.cells[i].width = w
        p = hdr5.cells[i].paragraphs[0]
        p.clear()
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8)
    style_header_row(hdr5)

    for idx, d in enumerate(detail, 1):
        row = table5.add_row()
        values = [
            str(idx), d["date"], d["staff"], d["project"],
            d["topic"], d["task"], d["subtask"], d["remarks"],
            f"{d['hours']:.2f}"
        ]
        for i, val in enumerate(values):
            p = row.cells[i].paragraphs[0]
            p.clear()
            run = p.add_run(val)
            run.font.size = Pt(8)

    # Total row for detail
    detail_total = sum(d["hours"] for d in detail)
    total_row5 = table5.add_row()
    for i in range(len(headers)):
        p = total_row5.cells[i].paragraphs[0]
        p.clear()
        run = p.add_run("TOTAL" if i == 7 else (f"{detail_total:.2f}" if i == 8 else ""))
        run.bold = True
        run.font.size = Pt(8)

    # ── SECTION 6: DECLARATION ───────────────────
    doc.add_page_break()
    doc.add_heading("6. Declaration", level=1)
    doc.add_paragraph(
        "This report has been prepared in accordance with MFRS 138 Intangible Assets and Section 34(7) "
        "of the Income Tax Act 1967 (Malaysia). The hours recorded represent qualifying R&D activities "
        "undertaken by Dallas Roboter Sdn Bhd for the financial year May 2025 to April 2026. "
        "All time entries have been extracted from the Company's Time Management System and represent "
        "actual hours logged by authorised personnel against approved R&D projects."
    )

    doc.add_paragraph()
    doc.add_paragraph("Prepared by: _______________________________")
    doc.add_paragraph("Name:        _______________________________")
    doc.add_paragraph("Designation: _______________________________")
    doc.add_paragraph(f"Date:        {date.today().strftime('%d %B %Y')}")

    # Footer
    add_footer(doc)

    return doc

# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────

def main():
    os.makedirs("output", exist_ok=True)

    print("Fetching data from database...")
    rows = fetch_data()

    print("Building summaries...")
    by_project, by_person, by_month, months_ordered, detail = build_summaries(rows)

    print("Generating Word report...")
    doc = build_report(by_project, by_person, by_month, months_ordered, detail)

    doc.save(OUTPUT_PATH)

    total_projects = len(by_project)
    total_staff    = len(by_person)
    total_hours    = sum(d["hours"] for d in detail)

    print("\nReport generated successfully!")
    print(f"Total projects : {total_projects}")
    print(f"Total staff    : {total_staff}")
    print(f"Total hours    : {total_hours:.2f}")
    print(f"Output saved to: {OUTPUT_PATH}")

if __name__ == "__main__":
    main()